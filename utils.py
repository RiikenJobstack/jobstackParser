import os
import io
import fitz  # PyMuPDF
import pdfplumber
import docx
from PIL import Image
import json
import hashlib
import pickle
from functools import lru_cache
from typing import Optional, Any
from dotenv import load_dotenv
import time
import threading
import json
import re
from typing import Optional, Dict, Any
import requests
import google.auth
import google.auth.transport.requests
# Import static prompts
from static_prompt import STATIC_RESUME_PARSER_PROMPT
# Cache utility functions - add these to your file

import hashlib
import pickle
from typing import Optional, Any
import os
# from resumeUtils import count_tokens, calculate_total_cost

load_dotenv()

# Initialize Google Generative AI client
import google.generativeai as genai
from google.oauth2 import service_account
import google.auth.transport.requests
import requests
# Initialize credentials at module level
GOOGLE_APPLICATION_CREDENTIALS  = os.environ.get("google-app-credentials")
# Set up credentials with proper scopes
credentials = service_account.Credentials.from_service_account_info(
    GOOGLE_APPLICATION_CREDENTIALS,
    scopes=["https://www.googleapis.com/auth/cloud-platform"]
)
genai.configure(credentials=credentials)

# Initialize Gemini model
gemini_model = genai.GenerativeModel('gemini-2.5-flash-lite')

# Alternative: Use Vertex AI REST API directly (like your Node.js code)
def call_gemini_api_direct(prompt: str, model: str = "gemini-2.5-flash-lite") -> str:
    """Call Gemini API directly using Vertex AI REST endpoint"""
    try:
        # Get access token
        auth_req = google.auth.transport.requests.Request()
        credentials.refresh(auth_req)
        access_token = credentials.token

        # Vertex AI endpoint
        endpoint = f"https://us-central1-aiplatform.googleapis.com/v1/projects/{GOOGLE_APPLICATION_CREDENTIALS['project_id']}/locations/us-central1/publishers/google/models/{model}:generateContent"

        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }

        data = {
            "contents": [
                {
                    "role": "USER",
                    "parts": [
                        {
                            "text": prompt
                        }
                    ]
                }
            ],
            "generationConfig": {
                "maxOutputTokens": 8192,
                "temperature": 0.2
            }
        }

        response = requests.post(endpoint, headers=headers, json=data, timeout=30)
        response.raise_for_status()

        result = response.json()

        # Extract text from response (same structure as your Node.js code)
        if (result.get("candidates") and
            result["candidates"][0].get("content") and
            result["candidates"][0]["content"].get("parts") and
            result["candidates"][0]["content"]["parts"][0].get("text")):

            return result["candidates"][0]["content"]["parts"][0]["text"]
        else:
            raise Exception("Invalid response structure from Gemini API")

    except Exception as e:
        print(f"Direct API call failed: {str(e)}")
        raise e

# Initialize Redis client (optional - graceful fallback if not available)
redis_client = None
try:
    import redis
    redis_client = redis.Redis(
        host=os.getenv('REDIS_HOST', 'localhost'),
        port=int(os.getenv('REDIS_PORT', 6379)),
        db=0,
        decode_responses=False,
        socket_connect_timeout=1,
        socket_timeout=1
    )
    redis_client.ping()  # Test connection
except Exception:
    # Redis not available or not installed - use only in-memory cache
    redis_client = None

# In-memory cache as fallback
_cache = {}
_max_cache_size = 1000
_cache_ttl = 86400 * 7  # 7 days

def _get_file_hash(content: bytes) -> str:
    """Generate hash of file content for cache keys"""
    return hashlib.md5(content).hexdigest()

def _get_text_hash(text: str) -> str:
    """Generate hash of text content for cache keys"""
    return hashlib.md5(text.encode('utf-8')).hexdigest()

def _get_from_cache(key: str) -> Optional[Any]:
    """Get item from cache (Redis first, then in-memory)"""
    try:
        if redis_client:
            cached = redis_client.get(key)
            if cached:
                return pickle.loads(cached)
    except Exception:
        pass

    return _cache.get(key)

def _set_cache(key: str, value: Any):
    """Set item in cache (both Redis and in-memory)"""
    try:
        if redis_client:
            redis_client.setex(key, _cache_ttl, pickle.dumps(value))
    except Exception:
        pass

    # Also store in memory with size limit
    if len(_cache) >= _max_cache_size:
        # Remove oldest 25% of items
        keys_to_remove = list(_cache.keys())[:_max_cache_size // 4]
        for k in keys_to_remove:
            del _cache[k]

    _cache[key] = value

# Cache the expensive OCR reader initialization
@lru_cache(maxsize=1)
def _get_ocr_reader():
    import easyocr
    return easyocr.Reader(['en'], gpu=False)

# Get cached reader instance
reader = _get_ocr_reader()

def extract_text_from_resume(filename: str, content: bytes) -> str:
    """Extract text with caching - same interface as original"""
    # Check cache first
    file_hash = _get_file_hash(content)
    cache_key = f"text_extract:{file_hash}"

    cached_text = _get_from_cache(cache_key)
    if cached_text is not None:
        return cached_text

    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(content)
    elif ext == ".docx":
        text = extract_text_from_docx(content)
    elif ext in [".png", ".jpg", ".jpeg"]:
        text = extract_text_from_image(content)
    elif ext in [".txt"]:
        text = extract_text_from_txt(content)
    else:
        return "Unsupported file format."

    # Cache the result
    _set_cache(cache_key, text)
    return text

def extract_text_from_pdf(content: bytes) -> str:
    """PDF text extraction with OCR fallback and caching (preserve hyperlinks)"""
    file_hash = _get_file_hash(content)

    text = ""
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            # Extract text
            page_text = page.extract_text() or ""
            text += page_text

            # Extract hyperlinks (if any)
            if hasattr(page, "hyperlinks") and page.hyperlinks:
                for link in page.hyperlinks:
                    uri = link.get("uri")
                    if uri:
                        # Add "label → link" format (or any structure you prefer)
                        text += f"\n[LINK] {uri}"

    if not text.strip():
        # Check OCR cache
        ocr_cache_key = f"pdf_ocr:{file_hash}"
        cached_ocr = _get_from_cache(ocr_cache_key)
        if cached_ocr is not None:
            return cached_ocr

        # Perform OCR and cache result
        text = extract_text_from_pdf_with_ocr(content)
        _set_cache(ocr_cache_key, text)

    return text


def extract_text_from_pdf_with_ocr(content: bytes) -> str:
    """OCR extraction for PDFs - same as original but using cached reader"""
    doc = fitz.open(stream=content, filetype="pdf")
    text = ""
    for page in doc:
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        import numpy as np
        img_np = np.array(img)
        result = reader.readtext(img_np, detail=0)
        text += "\n".join(result) + "\n"
    return text

def extract_text_from_docx(content: bytes) -> str:
    """Enhanced DOCX text extraction with hyperlinks in [LINK] format"""
    doc = docx.Document(io.BytesIO(content))
    text_parts = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        text_parts.append(_extract_paragraph_with_links(para))

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_texts = []
                for paragraph in cell.paragraphs:
                    para_text = _extract_paragraph_with_links(paragraph)
                    if para_text:
                        cell_texts.append(para_text)
                if cell_texts:
                    row_text.append(" ".join(cell_texts))
                else:
                    row_text.append("")
            if any(cell.strip() for cell in row_text):  # Only add non-empty rows
                text_parts.append("\t".join(row_text))

    return "\n".join([t for t in text_parts if t.strip()])


def _extract_paragraph_with_links(paragraph):
    """Extract paragraph text with hyperlinks in [LINK] format"""
    text_parts = []
    # Iterate through all child XML elements (runs, hyperlinks, etc.)
    for child in paragraph._element:
        # Case 1: it's a hyperlink element
        if child.tag.endswith("hyperlink"):
            r_id = child.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            if r_id and r_id in paragraph.part.rels:
                url = paragraph.part.rels[r_id].target_ref
                if url:
                    text_parts.append(f"[LINK] {url}")
        # Case 2: it's a normal run (text)
        elif child.tag.endswith("r"):
            texts = [node.text for node in child.findall(".//w:t", namespaces=child.nsmap) if node.text]
            if texts:
                text_parts.append("".join(texts))
    return " ".join(text_parts).strip()


# Alternative version that preserves more table structure
def extract_text_from_docx_structured(content: bytes) -> str:
    """DOCX extraction that better preserves table structure"""
    doc = docx.Document(io.BytesIO(content))
    text_parts = []
    
    # Process document elements in order
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            # Find corresponding paragraph object
            for para in doc.paragraphs:
                if para._element == element and para.text.strip():
                    text_parts.append(para.text)
                    break
                    
        elif element.tag.endswith('tbl'):  # Table
            # Find corresponding table object
            for table in doc.tables:
                if table._element == element:
                    text_parts.append("\n--- TABLE START ---")
                    
                    for i, row in enumerate(table.rows):
                        row_text = []
                        for cell in row.cells:
                            cell_content = []
                            for paragraph in cell.paragraphs:
                                if paragraph.text.strip():
                                    cell_content.append(paragraph.text.strip())
                            
                            cell_text = " ".join(cell_content) if cell_content else ""
                            row_text.append(cell_text)
                        
                        if any(cell.strip() for cell in row_text):
                            text_parts.append(f"Row {i+1}: " + " | ".join(row_text))
                    
                    text_parts.append("--- TABLE END ---\n")
                    break
    
    return "\n".join(text_parts)


# Simplified version if you just want all text
def extract_text_from_docx_simple(content: bytes) -> str:
    """Simple DOCX extraction - gets all text from paragraphs and tables"""
    doc = docx.Document(io.BytesIO(content))
    
    full_text = []
    
    # Get all text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Get all text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text.append(paragraph.text)
    
    # Filter out empty lines and join
    return "\n".join([text for text in full_text if text.strip()])

def extract_text_from_image(content: bytes) -> str:
    """Image OCR with caching - same interface as original"""
    file_hash = _get_file_hash(content)
    cache_key = f"image_ocr:{file_hash}"

    # Check cache first
    cached_text = _get_from_cache(cache_key)
    if cached_text is not None:
        return cached_text

    # Perform OCR using cached reader
    image = Image.open(io.BytesIO(content)).convert("RGB")
    import numpy as np
    img_np = np.array(image)
    result = reader.readtext(img_np, detail=0)
    text = "\n".join(result)

    # Cache the result
    _set_cache(cache_key, text)
    return text

def extract_text_from_txt(content: bytes) -> str:
    """Extract text from a TXT file (UTF-8 with fallback)"""
    try:
        # Try decoding as UTF-8 first
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        # Fallback to latin1 if UTF-8 fails
        text = content.decode("latin1")

    return text.strip()

def transform_text_to_resume_data(raw_text: str) -> dict:
    """Transform text to structured data with Gemini API caching - same interface"""
    # Check cache first
    text_hash = _get_text_hash(raw_text)
    cache_key = f"gemini_transform:{text_hash}"

    cached_result = _get_from_cache(cache_key)
    if cached_result is not None:
        return cached_result

    prompt = f"""
You are a resume parser. Extract resume data to this JSON structure:

{{
  "targetJobTitle": "",
  "targetJobDescription": "",
  "personalInfo": {{
    "fullName": "",
    "jobTitle": "",
    "email": "",
    "phone": "",
    "location": "",
    "summary": "",
    "linkedinUrl": "",
    "githubUrl": "",
    "profilePicture": null
  }},
  "sections": [
    {{
      "id": "experience-1",
      "type": "experience",
      "title": "Work Experience",
      "order": 0,
      "hidden": false,
      "items": [{{
        "id": "exp-0",
        "jobTitle": "",
        "company": "",
        "location": "",
        "startDate": "",
        "endDate": null,
        "currentPosition": false,
        "description": ""
      }}],
      "groups": [],
      "content": "",
      "state": {{}}
    }},
    {{
      "id": "education-1",
      "type": "education",
      "title": "Education",
      "order": 1,
      "hidden": false,
      "items": [{{
        "id": "edu-0",
        "degree": "",
        "institution": "",
        "location": "",
        "startDate": "",
        "endDate": "",
        "current": false,
        "gpa": "",
        "description": ""
      }}],
      "groups": [],
      "content": "",
      "state": {{}}
    }},
    {{
      "id": "skills-1",
      "type": "skills",
      "title": "Skills",
      "order": 2,
      "hidden": false,
      "format": "grouped",
      "items": ["string"], (all skills)
      "groups": [{{
        "id": "skill-0",
        "name": "", (skill category)
        "items": []
      }}],
      "content": "",
      "state": {{"categoryOrder": [], "viewMode": "categorized"}}
    }},
    {{
      "id": "projects-1",
      "type": "projects",
      "title": "Projects",
      "order": 3,
      "hidden": false,
      "items": [{{
        "id": "proj-0",
        "name": "",
        "description": "",
        "technologies": [],
        "url": "",
        "startDate": "",
        "endDate": "",
        "current": false
      }}],
      "groups": [],
      "content": "",
      "state": {{}}
    }},
    {{
      "id": "awards-1",
      "type": "awards",
      "title": "AWARDS & ACHIEVEMENTS",
      "order": 4,
      "hidden": false,
      "items": [{{
        "id": "award-0",
        "title": "",
        "organization": "",
        "date": "",
        "description": ""
      }}],
      "groups": [],
      "content": "",
      "state": {{}}
    }},
    {{
      "id": "certifications-1",
      "type": "certifications",
      "title": "CERTIFICATIONS",
      "order": 5,
      "hidden": false,
      "items": [{{
        "id": "cert-0",
        "name": "",
        "issuer": "",
        "date": "",
        "description": ""
      }}],
      "groups": [],
      "content": "",
      "state": {{}}
    }},
    {{
      "id": "languages-1",
      "type": "languages",
      "title": "LANGUAGES",
      "order": 6,
      "hidden": false,
      "items": [{{
        "id": "lang-0",
        "name": "",
        "proficiency": ""
      }}],
      "groups": [],
      "content": "",
      "state": {{}}
    }},
    {{
      "id": "volunteer-1",
      "type": "volunteer",
      "title": "VOLUNTEER EXPERIENCE",
      "order": 7,
      "hidden": false,
      "items": [{{
        "id": "vol-0",
        "role": "",
        "organization": "",
        "location": "",
        "startDate": "",
        "endDate": "",
        "description": ""
      }}],
      "groups": [],
      "content": "",
      "state": {{}}
    }},
    {{
      "id": "key-achievements-1",
      "type": "key-achievements",
      "title": "KEY ACHIEVEMENTS",
      "order": 8,
      "hidden": false,
      "items": [{{
        "id": "achieve-0",
        "name": "",
        "description": ""
      }}],
      "groups": [],
      "content": "",
      "state": {{}}
    }},
    {{
      "id": "interests-1",
      "type": "interests",
      "title": "INTERESTS",
      "order": 9,
      "hidden": false,
      "items": [{{
        "id": "int-0",
        "category": "",
        "items": "",
        "description": ""
      }}],
      "groups": [],
      "content": "",
      "state": {{}}
    }}
  ]
}}

RULES:
DO NOT include sections with empty/no data
Dates: YYYY-MM-DD format
Current positions: endDate=null, currentPosition/current=true
IDs: exp-0, edu-0, proj-0, cert-0, lang-0 etc.
Skills: categorize into groups (technical, frameworks, tools, soft)
URLs: Extract LinkedIn and GitHub URLs from resume text, social media sections, or contact info
LinkedIn URL formats: linkedin.com/in/username, www.linkedin.com/in/username
GitHub URL formats: github.com/username, www.github.com/username
Additional sections: Add any other sections (awards, volunteer, publications, etc.) as new section objects in sections array with appropriate id, type, title, order, hidden=false
Bullet points: return as array of strings
Paragraphs: return as single string inside array
Missing data: empty strings/arrays, null for dates
Return only valid JSON

Resume Text:
{raw_text}

Return only valid JSON.
"""

    try:
        # Use direct API call (same as your Node.js code)
        print("Calling Gemini API directly...")
        content = call_gemini_api_direct(prompt)

        # Debug: Log the response content for troubleshooting
        print(f"Gemini API Response Length: {len(content)}")
        print(f"Gemini API Response Preview: {content[:200]}...")

        if not content:
            error_result = {"error": "Gemini API returned empty content"}
            return error_result

        # Try to parse JSON from the response
        try:
            # First try to extract JSON from code blocks (like your Node.js code)
            import re
            json_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', content)
            if json_match:
                print("Found JSON in code block format")
                result = json.loads(json_match.group(1))
            else:
                print("Attempting to parse entire response as JSON")
                result = json.loads(content)

            # Cache successful result
            _set_cache(cache_key, result)
            return result

        except json.JSONDecodeError as e:
            print(f"JSON Parse Error: {str(e)}")
            print(f"Content that failed to parse: {content[:500]}...")
            error_result = {"error": f"Failed to parse JSON response: {str(e)}"}
            # Don't cache JSON parsing errors
            return error_result

    except Exception as e:
        print(f"Gemini API Exception: {str(e)}")
        error_result = {"error": f"Gemini API error: {str(e)}"}
        # Don't cache errors
        return error_result
    
def transform_text_to_resume_data_temp(raw_text: str) -> dict:
    print('raw_text',raw_text)
    """Transform text to structured data with Gemini API caching - same interface"""
    # Check cache first
    text_hash = _get_text_hash(raw_text)
    cache_key = f"gemini_transform:{text_hash}"

    cached_result = _get_from_cache(cache_key)
    if cached_result is not None:
        return cached_result

    prompt = f"""
You are an expert resume parser. Extract ONLY the content from the resume text below.
DO NOT include any structure, settings, metadata, IDs, or UI configuration.

Return a JSON object with this exact structure:

{{
  "success": true,
  "data": {{
    "content": {{
      "personalInfo": {{
        "fullName": "",
        "title": "",
        "email": "",
        "phone": "",
        "location": "",
        "linkedIn": "",
        "portfolio": "",
        "github": "",
        "customLinks": []
      }},
      "summary": {{
        "content": ""
      }},
      "experience": [
        {{
          "company": "",
          "position": "",
          "location": "",
          "startDate": null,
          "endDate": null,
          "current": false,
          "description": "",
          "achievements": [],
          "technologies": [],
          "employmentType": "Full-time",
          "remote": false
        }}
      ],
      "projects": [
        {{
          "title": "",
          "role": "",
          "type": "personal",
          "client": "",
          "startDate": null,
          "endDate": null,
          "current": false,
          "url": "",
          "github": "",
          "description": "",
          "technologies": [],
          "achievements": [],
          "teamSize": "",
          "impact": ""
        }}
      ],
      "education": [
        {{
          "institution": "",
          "degree": "",
          "field": "",
          "location": "",
          "startDate": null,
          "endDate": null,
          "current": false,
          "gpa": "",
          "achievements": [],
          "courses": [],
          "honors": [],
          "online": false
        }}
      ],
      "skills": {{
        "extracted": []
      }},
      "certifications": [
        {{
          "name": "",
          "issuer": "",
          "issueDate": null,
          "expiryDate": null,
          "credentialId": "",
          "url": "",
          "skills": []
        }}
      ],
      "awards": [
        {{
          "title": "",
          "issuer": "",
          "date": null,
          "description": "",
          "category": "",
          "amount": ""
        }}
      ],
      "languages": [
        {{
          "language": "",
          "proficiency": "",
          "certification": ""
        }}
      ],
      "volunteering": [
        {{
          "organization": "",
          "role": "",
          "cause": "",
          "location": "",
          "startDate": null,
          "endDate": null,
          "current": false,
          "description": "",
          "impact": "",
          "hoursPerWeek": null
        }}
      ],
      "publications": [
        {{
          "title": "",
          "authors": [],
          "publisher": "",
          "date": null,
          "url": "",
          "description": "",
          "type": "journal",
          "doi": "",
          "conference": "",
          "citations": null,
          "journal": ""
        }}
      ]
    }},
    "parseMetadata": {{
      "confidence": 0.0,
      "parseTime": 0.0,
      "detectedSections": [],
      "missingSections": [],
      "sectionConfidence": {{
        "personalInfo": 0.0,
        "experience": 0.0,
        "education": 0.0,
        "skills": 0.0,
        "projects": 0.0
      }},
      "warnings": [
        {{
          "type": "",
          "message": "",
          "section": "",
          "field": "",
          "suggestion": ""
        }}
      ],
      "suggestions": [
        {{
          "section": "",
          "type": "",
          "message": "",
          "priority": "",
          "example": ""
        }}
      ],
      "extractedKeywords": [],
      "industryDetected": "",
      "experienceLevel": "",
      "totalExperienceYears": null,
      "educationLevel": "",
      "atsKeywords": {{
        "technical": [],
        "soft": [],
        "industry": [],
        "certifications": []
      }},
      "stats": {{
        "totalWords": 0,
        "bulletPoints": 0,
        "quantifiedAchievements": 0,
        "actionVerbs": 0,
        "uniqueSkills": 0
      }}
    }}
  }}
}}

PARSING RULES:

1. CONTENT EXTRACTION:
   - Extract ONLY actual content from the resume
   - DO NOT generate IDs, orders, enabled states, or any UI configuration
   - DO NOT include template, theme, or layout information
   - Return null for missing fields, empty arrays for missing lists
   - Only include sections that have actual content
   
   HANDLING [LINK] FORMAT:
   - The resume text may contain links formatted as: [LINK] followed by URL
   - These need to be parsed and mapped to appropriate fields
   - Examples:
     * [LINK] mailto:email@example.com → email field
     * [LINK] tel:+91-9999999999 → phone field
     * [LINK] https://linkedin.com/in/username → linkedIn field
     * [LINK] https://github.com/username → github field
     * [LINK] https://projectname.com → project url if near project context
   - Always clean URLs by removing protocols and prefixes

2. PERSONAL INFO:
   - fullName: Complete name as found in resume
   - title: Professional title or role they're targeting (e.g., "Senior Software Engineer")
   - email: Email address (validate format)
   - phone: Phone number with country code if available and remove any special characters other than number like '-','*' etc. EX - if phone is +91-8709910113 then give only +918709910113
   - location: "City, State" or "City, Country" format
   - linkedIn: LinkedIn URL without https:// (e.g., "linkedin.com/in/username")
   - portfolio: Portfolio/personal website URL without https://
   - github: GitHub URL without https:// (e.g., "github.com/username")
   - customLinks: Array of {{"label": "Link Text", "url": "domain.com/path"}} for other links
   
   SPECIAL LINK HANDLING:
   - Resume may contain links in format: [LINK] https://example.com or [LINK] mailto:email@example.com
   - Map [LINK] entries correctly:
     * [LINK] mailto:xxx → extract email address (remove mailto:)
     * [LINK] tel:xxx → extract phone number (remove tel:) and remove any special characters other than number like '-','*' etc. EX - if phone is +91-8709910113 then give only +918709910113
     * [LINK] containing linkedin.com → linkedIn field (remove https://)
     * [LINK] containing github.com → github field (remove https://)
     * Other [LINK] URLs → Check if mentioned near a project name, then add to that project's url field
     * Remaining links → portfolio field if it looks like a portfolio, otherwise ignore
   - Clean all URLs by removing https://, http://, www.

3. DATES:
   - Always use YYYY-MM-DD format for all dates
   - For current positions: endDate = null, current = true
   - For date ranges like "Jan 2020 - Present": startDate = "2020-01-01", endDate = null, current = true
   - For single dates, use the first day of the period
   - If only year: "2020" becomes "2020-01-01"
   - If month and year: "Jan 2020" becomes "2020-01-01"
   - If date is unclear or missing: use null

4. EXPERIENCE & EDUCATION DETAILS:
   
   EXPERIENCE fields:
   - company: Company/Organization name
   - position: Job title/role
   - location: Job location (City, State/Country)
   - startDate/endDate: Employment period (YYYY-MM-DD)
   - current: true if currently employed there
   - description: Overall role description or main responsibilities (paragraph text)
   - achievements: Array of achievements/accomplishments (clean bullet points)
   - technologies: Technologies/tools used in this role
   - employmentType: "Full-time" | "Part-time" | "Contract" | "Freelance" | "Internship" | "Volunteer"
   - remote: true if remote position
   
   EDUCATION fields:
   - institution: School/University name
   - degree: Degree type (e.g., "Bachelor of Science", "Master of Arts")
   - field: Field of study (e.g., "Computer Science", "Business Administration")
   - location: School location
   - startDate/endDate: Study period (YYYY-MM-DD)
   - current: true if currently studying
   - gpa: GPA if mentioned (e.g., "3.8/4.0")
   - achievements: Array of honors, awards, achievements
   - courses: Array of relevant coursework
   - honors: Array of honors (Dean's list, cum laude, etc.)
   - online: true if online degree program

5. SKILLS:
   - Extract all skills mentioned anywhere in the resume
   - Return as simple flat array in "extracted" field
   - Include programming languages, frameworks, tools, soft skills, everything
   - Remove duplicates (case-insensitive)
   - Preserve original capitalization (e.g., "JavaScript" not "javascript")
   - Order by relevance/frequency of mention if possible

6. PROJECTS (include only if found):
   - name: Project name
   - role: Person's role in the project (e.g., "Lead Developer", "Contributor")
   - type: "personal" | "professional" | "academic" | "open-source" | "freelance"
   - client: Client name for freelance projects
   - startDate/endDate: Project timeline in YYYY-MM-DD format
   - current: true if ongoing project
   - url: Project demo/live URL without https://
   - github: GitHub repository URL without https://
   - description: Brief project description
   - technologies: Array of technologies/tools used
   - achievements: Array of specific outcomes/metrics
   - teamSize: Team size if mentioned (e.g., "5-10 people")
   - impact: Business impact or metrics achieved
   
   PROJECT LINK MAPPING:
   - If a [LINK] URL appears near a project name/description, assign it to that project's url field
   - If the URL contains github.com, use it for the github field instead
   - Example: "CitiTour [LINK] https://example.com/" → url: "example.com"
   - Clean URLs by removing https://, http://, www., and trailing slashes

7. ADDITIONAL SECTIONS (include only if content exists):
   
   CERTIFICATIONS:
   - name: Certification name
   - issuer: Issuing organization
   - issueDate: Date obtained (YYYY-MM-DD)
   - expiryDate: Expiry date if applicable
   - credentialId: Credential/License ID if mentioned
   - url: Verification URL without https://
   - skills: Array of skills related to this certification
   
   AWARDS:
   - title: Award/Honor title
   - issuer: Organization that gave the award
   - date: Date received (YYYY-MM-DD)
   - description: Brief description
   - category: Type (e.g., "Academic Excellence", "Professional Achievement")
   - amount: For scholarships or monetary awards
   
   LANGUAGES:
   - language: Language name
   - proficiency: "Native" | "Fluent" | "Advanced" | "Intermediate" | "Basic"
   - certification: Any language certification (e.g., "TOEFL 110")
   
   VOLUNTEERING:
   - organization: Organization name
   - role: Volunteer position
   - cause: Cause or focus area
   - location: Location
   - startDate/endDate: Timeline (YYYY-MM-DD)
   - current: true if still volunteering
   - description: Activities description
   - impact: Impact or achievements
   - skills: Skills used or developed
   - hours: Total hours as number if mentioned
   
   PUBLICATIONS:
   - title: Publication title
   - authors: Array of author names
   - publisher: Publisher or journal name
   - date: Publication date (YYYY-MM-DD)
   - url: Publication URL without https://
   - description: Abstract or description
   - type: "journal" | "conference" | "book" | "chapter" | "thesis" | "patent" | "other"
   - doi: DOI if available
   - conference: Conference name if applicable
   - citations: Number of citations if mentioned
   - journal: Journal name if applicable

8. SECTIONS TO DETECT:
   - Only include sections that actually have content.
        IMPORTANT: Only include sections in the output that contain actual data. Do NOT include
        empty sections.
        - If a section has no data (empty array or null values), completely OMIT that section
        from the response
        - Examples:
          * If no languages found → DO NOT include "languages" key at all
          * If no publications → DO NOT include "publications" key
          * If no volunteering → DO NOT include "volunteering" key
          * If no awards → DO NOT include "awards" key
          * If no certifications → DO NOT include "certifications" key

        Good example:
        {{
          "personalInfo": {...},
          "experience": [...],
          "skills": {...}
          // No empty sections included
        }}

        Bad example:
        {{
          "personalInfo": {...},
          "experience": [...],
          "skills": {...},
          "languages": [],  // ❌ Don't include empty arrays
          "awards": [],      // ❌ Don't include if no data
          "publications": [] // ❌ Omit completely
        }}
   - Don't create empty sections
   - Common section variations to recognize:
     * Work/Professional Experience → experience
     * Academic Background/Education → education
     * Technical Skills/Core Competencies → skills
     * Projects/Personal Projects → projects
     * Honors/Awards/Achievements → awards
     * Volunteer/Community Service → volunteer
     * Publications/Research → publications
     * Certifications/Licenses → certifications
     * Languages/Language Proficiency → languages

8. URLS:
   - Remove "https://", "http://", "www." from all URLs
   - Store as clean domain paths (e.g., "linkedin.com/in/username")
   - For portfolio/personal sites, keep full domain
   - If URL is invalid or malformed, set to null

9. CONFIDENCE & METADATA:
   - confidence: 0.0 to 1.0 overall parsing confidence
   - parseTime: Will be filled by system (leave as 0.0)
   - detectedSections: Array of section names found (e.g., ["personalInfo", "experience", "education", "skills"])
   - missingSections: Important sections that are missing (e.g., ["summary", "projects"])
   - sectionConfidence: Confidence score (0.0-1.0) for each major section
   - extractedKeywords: Top 20 important technical/professional keywords found
   - atsKeywords: Categorized keywords for ATS optimization:
     * technical: Programming languages, frameworks, databases (e.g., "Python", "React", "PostgreSQL")
     * soft: Soft skills (e.g., "Leadership", "Communication", "Problem-solving")
     * industry: Industry-specific terms (e.g., "Agile", "CI/CD", "Machine Learning")
     * certifications: Certification keywords (e.g., "AWS Certified", "PMP", "CPA")
   - industryDetected: Detected industry (e.g., "Software Engineering", "Marketing", "Healthcare")
   - experienceLevel: Must be exactly one of: "Entry" | "Mid" | "Senior" | "Executive" (NOT "Mid-level", "Entry-level", etc.)
     * "Entry": 0-2 years experience
     * "Mid": 2-5 years experience  
     * "Senior": 5-10 years experience
     * "Executive": 10+ years experience
   - totalExperienceYears: Calculated total years of experience as a number
   - educationLevel: Must be exactly one of: "High School" | "Associate" | "Bachelor's" | "Master's" | "PhD" | "Professional"
     * Use "Bachelor's" NOT "Bachelor's Degree"
     * Use "Master's" NOT "Master's Degree"
     * Use "Associate" NOT "Associate's Degree"
   - atsScore: 0-100 score for ATS compatibility based on:
     * Contact info completeness (20 points)
     * Professional summary (15 points)
     * Work experience (25 points)
     * Education (15 points)
     * Skills section (15 points)
     * Keywords and formatting (10 points)
   - stats:
     * totalWords: Total word count in resume
     * bulletPoints: Number of bullet points found
     * quantifiedAchievements: Number of achievements with metrics/numbers
     * actionVerbs: Count of strong action verbs used
     * uniqueSkills: Number of unique skills identified

10. WARNINGS TO GENERATE (with structure):
    Each warning should have:
    - type: "missing_field" | "date_format" | "low_confidence" | "formatting_issue" | "data_quality"
    - message: Human-readable warning (e.g., "Email address not found")
    - section: Which section is affected (e.g., "personalInfo", "experience")
    - field: Specific field if applicable (e.g., "email", "startDate")
    - severity: "low" (minor issue) | "medium" (should fix) | "high" (critical issue)
    
    Examples:
    - Missing email/phone (high severity)
    - Ambiguous dates like "2020-Present" without specific months (low severity)
    - Experience without descriptions (medium severity)
    - Skills section missing (medium severity)

11. SUGGESTIONS TO GENERATE (with structure):
    Each suggestion should have:
    - section: Which section to improve (e.g., "experience", "skills")
    - type: "add_metrics" | "expand_content" | "add_keywords" | "improve_formatting" | "add_section"
    - message: Actionable suggestion (e.g., "Add quantifiable metrics to your achievements")
    - priority: "low" (nice to have) | "medium" (recommended) | "high" (critical for ATS)
    - example: Concrete example of improvement (e.g., "Instead of 'Managed team', say 'Managed team of 5 engineers'")
    
    Examples:
    - Add metrics to achievements (high priority)
    - Include industry keywords (high priority)
    - Add professional summary (medium priority)
    - Expand project descriptions (low priority)

12. TEXT PROCESSING:
    - Preserve bullet points as array items (clean text only)
    - Keep paragraph text as single strings
    - Clean up extra whitespace
    - Fix common OCR errors if detected
    - Preserve emphasis if possible (but as plain text)

13. CUSTOM SECTIONS:
    - Any section not matching standard types goes in customSections
    - Preserve original section title
    - type: "text" for paragraph content, "list" for bullet points

14. QUALITY CHECKS:
    - Ensure valid JSON output
    - No trailing commas
    - Proper null values (not "null" strings)
    - Valid boolean values (true/false, not "true"/"false")
    - Empty arrays [] for lists with no items
    - null for missing single values

IMPORTANT:
- Return ONLY the JSON structure
- No explanations or text outside the JSON
- Core sections (personalInfo, summary, experience, education, skills) MUST always be included even if empty
- Optional sections (projects, certifications, awards, languages, volunteering, publications) MUST BE COMPLETELY EXCLUDED if they have no content - DO NOT include them as empty arrays
- For example: If no certifications found, DO NOT include "certifications": [], just omit the field entirely

EXAMPLE OF EXPECTED METADATA OUTPUT:
{{
  "parseMetadata": {{
    "confidence": 0.92,
    "parseTime": 0.0,
    "detectedSections": ["personalInfo", "experience", "education", "skills", "projects"],
    "missingSections": ["summary", "certifications"],
    "warnings": [
      {{
        "type": "missing_field",
        "message": "LinkedIn profile URL not found",
        "section": "personalInfo",
        "field": "linkedIn",
        "severity": "low"
      }}
    ],
    "suggestions": [
      {{
        "section": "experience",
        "type": "add_metrics",
        "message": "Add quantifiable achievements to your work experience",
        "priority": "high",
        "example": "Increased sales by 25% in Q3 2023"
      }},
      {{
        "section": "summary",
        "type": "add_section",
        "message": "Add a professional summary to improve ATS score",
        "priority": "medium",
        "example": "Results-driven software engineer with 5+ years of experience..."
      }}
    ],
    "extractedKeywords": ["React", "Node.js", "AWS", "Python", "Docker", "Agile", "CI/CD", "JavaScript", "MongoDB", "REST API"],
    "industryDetected": "Software Engineering",
    "experienceLevel": "Senior",
    "totalExperienceYears": 7.5,
    "educationLevel": "Bachelor's",
    "atsKeywords": {{
      "technical": ["React", "Node.js", "Python", "JavaScript", "MongoDB"],
      "soft": ["Leadership", "Communication", "Team Collaboration"],
      "industry": ["Agile", "CI/CD", "REST API", "Microservices"],
      "certifications": []
    }},
    "atsScore": 78,
    "sectionConfidence": {{
      "personalInfo": 0.95,
      "experience": 0.90,
      "education": 0.88,
      "skills": 0.93,
      "projects": 0.85
    }},
    "stats": {{
      "totalWords": 450,
      "bulletPoints": 12,
      "quantifiedAchievements": 3,
      "actionVerbs": 15,
      "uniqueSkills": 24
    }}
  }}
}}

Resume Text to Parse:
{raw_text}
"""
    
    try:
        # Use direct API call (same as your Node.js code)
        print("Calling Gemini API directly...")
        content = call_gemini_api_direct(prompt)

        # Debug: Log the response content for troubleshooting
        print(f"Gemini API Response Length: {len(content)}")
        print(f"Gemini API Response Preview: {content[:200]}...")

        if not content:
            error_result = {"error": "Gemini API returned empty content"}
            return error_result

        # Try to parse JSON from the response
        try:
            # First try to extract JSON from code blocks (like your Node.js code)
            import re
            json_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', content)
            if json_match:
                print("Found JSON in code block format")
                result = json.loads(json_match.group(1))
            else:
                print("Attempting to parse entire response as JSON")
                result = json.loads(content)

            # Cache successful result
            _set_cache(cache_key, result)
            return result

        except json.JSONDecodeError as e:
            print(f"JSON Parse Error: {str(e)}")
            print(f"Content that failed to parse: {content[:500]}...")
            error_result = {"error": f"Failed to parse JSON response: {str(e)}"}
            # Don't cache JSON parsing errors
            return error_result

    except Exception as e:
        print(f"Gemini API Exception: {str(e)}")
        error_result = {"error": f"Gemini API error: {str(e)}"}
        # Don't cache errors
        return error_result


def parse_resume(filename: str, content: bytes) -> dict:
    """Main parsing function with full pipeline caching - same interface as original"""
    # Check for complete cached result first
    file_hash = _get_file_hash(content)
    cache_key = f"full_parse:{file_hash}"

    cached_result = _get_from_cache(cache_key)
    if cached_result is not None:
        return cached_result

    # Process normally with individual step caching
    raw_text = extract_text_from_resume(filename, content)
    structured_data = transform_text_to_resume_data(raw_text)

    # Cache the complete result
    _set_cache(cache_key, structured_data)
    return structured_data

def parse_resume_temp(filename: str, content: bytes) -> dict:
    """Main parsing function with full pipeline caching - same interface as original"""
    # Check for complete cached result first
    file_hash = _get_file_hash(content)
    cache_key = f"full_parse:{file_hash}"

    cached_result = _get_from_cache(cache_key)
    if cached_result is not None:
        return cached_result

    # Process normally with individual step caching
    raw_text = None
    if filename is 'resume.txt':
        raw_text = content
    else:
        raw_text = extract_text_from_resume(filename, content)
    structured_data = transform_text_to_resume_data_temp(raw_text)
    # Cache the complete result
    _set_cache(cache_key, structured_data)
    return structured_data

# Optional: Cache management functions for monitoring
def get_cache_stats():
    """Get cache statistics for monitoring"""
    stats = {
        "in_memory_size": len(_cache),
        "redis_available": redis_client is not None
    }

    if redis_client:
        try:
            info = redis_client.info()
            stats["redis_used_memory"] = info.get("used_memory_human", "N/A")
            stats["redis_keys"] = redis_client.dbsize()
        except Exception:
            stats["redis_error"] = "Could not get Redis stats"

    return stats

def clear_cache():
    """Clear all caches"""
    global _cache
    _cache.clear()

    if redis_client:
        try:
            redis_client.flushdb()
        except Exception:
            pass

# Initialize Redis client (optional - graceful fallback if not available)
redis_client = None
try:
    import redis
    redis_client = redis.Redis(
        host=os.getenv('REDIS_HOST', 'localhost'),
        port=int(os.getenv('REDIS_PORT', 6379)),
        db=0,
        decode_responses=False,
        socket_connect_timeout=1,
        socket_timeout=1
    )
    redis_client.ping()  # Test connection
except Exception:
    # Redis not available or not installed - use only in-memory cache
    redis_client = None

# In-memory cache as fallback
_cache = {}
_max_cache_size = 1000
_cache_ttl = 86400 * 7  # 7 days

def _get_file_hash(content: bytes) -> str:
    """Generate hash of file content for cache keys"""
    return hashlib.md5(content).hexdigest()

def _get_text_hash(text: str) -> str:
    """Generate hash of text content for cache keys"""
    return hashlib.md5(text.encode('utf-8')).hexdigest()

def _get_from_cache(key: str) -> Optional[Any]:
    """Get item from cache (Redis first, then in-memory)"""
    try:
        if redis_client:
            cached = redis_client.get(key)
            if cached:
                return pickle.loads(cached)
    except Exception:
        pass

    return _cache.get(key)

def _set_cache(key: str, value: Any):
    """Set item in cache (both Redis and in-memory)"""
    try:
        if redis_client:
            redis_client.setex(key, _cache_ttl, pickle.dumps(value))
    except Exception:
        pass

    # Also store in memory with size limit
    if len(_cache) >= _max_cache_size:
        # Remove oldest 25% of items
        keys_to_remove = list(_cache.keys())[:_max_cache_size // 4]
        for k in keys_to_remove:
            del _cache[k]

    _cache[key] = value

# Optional: Cache management functions for monitoring
def get_cache_stats():
    """Get cache statistics for monitoring"""
    stats = {
        "in_memory_size": len(_cache),
        "redis_available": redis_client is not None
    }

    if redis_client:
        try:
            info = redis_client.info()
            stats["redis_used_memory"] = info.get("used_memory_human", "N/A")
            stats["redis_keys"] = redis_client.dbsize()
        except Exception:
            stats["redis_error"] = "Could not get Redis stats"

    return stats

def clear_cache():
    """Clear all caches"""
    global _cache
    _cache.clear()

    if redis_client:
        try:
            redis_client.flushdb()
        except Exception:
            pass


# Set up credentials with proper scopes
credentials = service_account.Credentials.from_service_account_info(
    GOOGLE_APPLICATION_CREDENTIALS,
    scopes=["https://www.googleapis.com/auth/cloud-platform"]
)

# Global cache variables for prompt caching
resume_cache_id = None
resume_cache_expiry = None
resume_cache_lock = threading.Lock()

def create_prompt_cache(static_instructions: str, model: str) -> str:
    """Create cache entry for static instructions"""
    try:
        # Get access token
        auth_req = google.auth.transport.requests.Request()
        credentials.refresh(auth_req)
        access_token = credentials.token

        endpoint = f"https://us-central1-aiplatform.googleapis.com/v1/projects/{GOOGLE_APPLICATION_CREDENTIALS['project_id']}/locations/us-central1/cachedContents"

        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }

        data = {
            "model": f"projects/{GOOGLE_APPLICATION_CREDENTIALS['project_id']}/locations/us-central1/publishers/google/models/{model}",
            "contents": [
                {
                    "role": "user",
                    "parts": [{"text": static_instructions}]
                }
            ],
            "ttl": "86400s"  # 24 hours
        }

        response = requests.post(endpoint, headers=headers, json=data, timeout=30)
        response.raise_for_status()

        result = response.json()
        return result["name"]

    except Exception as e:
        print(f"Cache creation failed: {str(e)}")
        raise e

def init_prompt_cache(model: str) -> str:
    """Initialize and manage prompt cache with thread safety"""
    global resume_cache_id, resume_cache_expiry

    now = time.time() * 1000  # Convert to milliseconds

    if resume_cache_id and resume_cache_expiry and now < resume_cache_expiry:
        return resume_cache_id
        
    with resume_cache_lock:
        # Double-check after acquiring lock
        if resume_cache_id and resume_cache_expiry and now < resume_cache_expiry:
            return resume_cache_id
        
        print("Creating resume parser prompt cache...")
        resume_cache_id = create_prompt_cache(STATIC_RESUME_PARSER_PROMPT, model)
        resume_cache_expiry = time.time() * 1000 + (23 * 60 * 60 * 1000)  # 23 hours
        print(f"Resume Parser Cache created: {resume_cache_id}")
        return resume_cache_id

def build_dynamic_prompt(resume_text: str) -> str:
    """Build dynamic prompt that references cached instructions"""
    return f"""Use the cached instructions.

<<<RESUME_START>>>
{resume_text}
<<<RESUME_END>>>

RESPONSE FORMAT:
Return ONLY the JSON object as specified in the cached schema. No markdown, no explanations."""

def count_tokens(prompt: str, model: str) -> int:
    """Count tokens in the prompt"""
    try:
        # Get access token
        auth_req = google.auth.transport.requests.Request()
        credentials.refresh(auth_req)
        access_token = credentials.token

        endpoint = f"https://us-central1-aiplatform.googleapis.com/v1/projects/{GOOGLE_APPLICATION_CREDENTIALS['project_id']}/locations/us-central1/publishers/google/models/{model}:countTokens"

        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }

        data = {
            "contents": [
                {
                    "role": "USER",
                    "parts": [{"text": prompt}]
                }
            ]
        }

        response = requests.post(endpoint, headers=headers, json=data, timeout=30)
        response.raise_for_status()

        result = response.json()
        return result.get("totalTokens", 0)

    except Exception as e:
        print(f"Token counting failed: {str(e)}")
        return 0

def calculate_total_cost(input_tokens: int, output_tokens: int) -> float:
    """Calculate total cost based on token usage"""
    input_cost = input_tokens * 0.25 / 1_000_000
    output_cost = output_tokens * 0.75 / 1_000_000
    return input_cost + output_cost

def call_gemini_with_cache_and_retry(dynamic_prompt: str, model: str, max_retries: int = 2) -> str:
    """Call Gemini API with caching and retry logic"""
    global resume_cache_id, resume_cache_expiry
    
    retries = 0

    while retries < max_retries:
        try:
            cache_id = init_prompt_cache(model)
            
            # Get access token
            auth_req = google.auth.transport.requests.Request()
            credentials.refresh(auth_req)
            access_token = credentials.token

            endpoint = f"https://us-central1-aiplatform.googleapis.com/v1/projects/{GOOGLE_APPLICATION_CREDENTIALS['project_id']}/locations/us-central1/publishers/google/models/{model}:generateContent"

            headers = {
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json"
            }

            data = {
                "cachedContent": cache_id,
                "contents": [
                    {
                        "role": "user",
                        "parts": [{"text": dynamic_prompt}]
                    }
                ],
                "generationConfig": {
                    "maxOutputTokens": 8000,
                    "temperature": 0.1,
                    "topP": 0.8,
                    "topK": 40
                }
            }

            response = requests.post(endpoint, headers=headers, json=data, timeout=30)
            response.raise_for_status()

            result = response.json()

            # Extract text from response
            if (result.get("candidates") and
                result["candidates"][0].get("content") and
                result["candidates"][0]["content"].get("parts") and
                result["candidates"][0]["content"]["parts"][0].get("text")):

                return result["candidates"][0]["content"]["parts"][0]["text"]
            else:
                raise Exception("Invalid response structure from Gemini API")

        except Exception as err:
            if "Cache content" in str(err) and retries < max_retries:
                print(f"Gemini cache expired, retrying attempt {retries + 1}")
                # Force refresh local cache
                resume_cache_id = None
                resume_cache_expiry = None
                retries += 1
                continue
            raise err


def parse_gemini_response(content: str) -> Dict[str, Any]:
    """
    Parse Gemini response with improved regex patterns and error handling
    """
    if not content or not isinstance(content, str):
        raise ValueError("Invalid content provided")

    print(f"Parsing content (length: {len(content)})")
    
    # Strategy 1: Try direct JSON parse
    try:
        result = json.loads(content.strip())
        print("✅ Strategy 1: Direct parse successful")
        return result
    except json.JSONDecodeError as e:
        print(f"❌ Strategy 1 failed: {str(e)}")

    # Strategy 2: Extract from code blocks with improved patterns
    print("🔄 Strategy 2: Trying code block extraction...")
    
    # More comprehensive code block patterns
    code_block_patterns = [
        # Standard markdown code blocks
        r'```json\s*\n?(.*?)\n?```',
        r'```\s*\n?(.*?)\n?```',
        # Code blocks without newlines
        r'```json\s*(.*?)\s*```',
        r'```\s*(.*?)\s*```',
        # Single backticks with JSON
        r'`\s*(\{.*?\})\s*`',
        # JSON wrapped in any amount of backticks
        r'`+\s*(\{[\s\S]*?\})\s*`+',
    ]

    for i, pattern in enumerate(code_block_patterns):
        print(f"   Trying pattern {i+1}: {pattern}")
        try:
            # Use DOTALL flag to match newlines in the content
            matches = re.findall(pattern, content, re.IGNORECASE | re.DOTALL)
            if matches:
                print(f"   Found {len(matches)} matches")
                for j, match in enumerate(matches):
                    try:
                        print(f"   Processing match {j+1} (length: {len(match)})")
                        # Clean the match more thoroughly
                        cleaned = match.strip()
                        
                        # Remove any remaining backticks or markdown artifacts
                        cleaned = re.sub(r'^`+|`+$', '', cleaned)
                        cleaned = re.sub(r'^json\s*', '', cleaned, flags=re.IGNORECASE)
                        cleaned = cleaned.strip()
                        
                        print(f"   Cleaned match preview: {cleaned[:200]}...")
                        result = json.loads(cleaned)
                        print("✅ Strategy 2: Code block extraction successful")
                        return result
                    except json.JSONDecodeError as e:
                        print(f"   Match {j+1} failed: {str(e)}")
                        continue
            else:
                print(f"   No matches found for pattern {i+1}")
        except Exception as e:
            print(f"   Pattern {i+1} error: {str(e)}")
            continue

    # Strategy 3: Find JSON-like content with better regex
    print("🔄 Strategy 3: Trying JSON pattern extraction...")
    try:
        # Look for JSON object that starts with { and ends with }
        # This pattern handles nested braces correctly
        json_pattern = r'\{(?:[^{}]|(?:\{[^{}]*\})*)*\}'
        matches = re.findall(json_pattern, content, re.DOTALL)
        
        if matches:
            # Try the longest match first (most likely to be complete)
            matches.sort(key=len, reverse=True)
            for i, match in enumerate(matches[:3]):  # Try top 3 matches
                try:
                    print(f"   Trying JSON match {i+1} (length: {len(match)})")
                    result = json.loads(match)
                    print("✅ Strategy 3: JSON pattern extraction successful")
                    return result
                except json.JSONDecodeError as e:
                    print(f"   JSON match {i+1} failed: {str(e)}")
                    continue
        else:
            print("   No JSON patterns found")
    except Exception as e:
        print(f"   Strategy 3 error: {str(e)}")

    # Strategy 4: Advanced JSON cleaning and fixing
    print("🔄 Strategy 4: Advanced JSON fixing...")
    try:
        # Start with the original content
        fixed_content = content
        
        # Remove markdown code block markers
        fixed_content = re.sub(r'```(?:json)?\s*\n?', '', fixed_content, flags=re.IGNORECASE)
        fixed_content = re.sub(r'\n?```\s*$', '', fixed_content)
        
        # Remove any leading/trailing backticks
        fixed_content = re.sub(r'^`+|`+$', '', fixed_content.strip())
        
        # Remove "json" keyword if it appears at the start
        fixed_content = re.sub(r'^\s*json\s*', '', fixed_content, flags=re.IGNORECASE)
        
        # Fix common JSON issues
        fixed_content = re.sub(r',(\s*[}\]])', r'\1', fixed_content)  # Remove trailing commas
        fixed_content = re.sub(r'([{,]\s*)(\w+):', r'\1"\2":', fixed_content)  # Quote unquoted keys
        fixed_content = re.sub(r":\s*'([^']*)'", r': "\1"', fixed_content)  # Replace single quotes
        
        # Clean up whitespace
        fixed_content = fixed_content.strip()
        
        print(f"   Fixed content preview: {fixed_content[:200]}...")
        result = json.loads(fixed_content)
        print("✅ Strategy 4: Advanced JSON fixing successful")
        return result
    except json.JSONDecodeError as e:
        print(f"   Strategy 4 failed: {str(e)}")
    except Exception as e:
        print(f"   Strategy 4 error: {str(e)}")

    # Strategy 5: Extract JSON from specific patterns in your response
    print("🔄 Strategy 5: Trying response-specific extraction...")
    try:
        # Look for the specific pattern in your error: ```json\n{ ... }
        pattern = r'```json\s*\n(\{[\s\S]*)'
        match = re.search(pattern, content, re.IGNORECASE)
        if match:
            json_content = match.group(1)
            # Find the end of the JSON by counting braces
            brace_count = 0
            end_pos = 0
            in_string = False
            escape_next = False
            
            for i, char in enumerate(json_content):
                if escape_next:
                    escape_next = False
                    continue
                    
                if char == '\\':
                    escape_next = True
                    continue
                    
                if char == '"' and not escape_next:
                    in_string = not in_string
                    continue
                    
                if not in_string:
                    if char == '{':
                        brace_count += 1
                    elif char == '}':
                        brace_count -= 1
                        if brace_count == 0:
                            end_pos = i + 1
                            break
            
            if end_pos > 0:
                complete_json = json_content[:end_pos]
                print(f"   Extracted complete JSON (length: {len(complete_json)})")
                result = json.loads(complete_json)
                print("✅ Strategy 5: Response-specific extraction successful")
                return result
    except Exception as e:
        print(f"   Strategy 5 error: {str(e)}")

    # If all strategies fail, provide detailed error information
    content_preview = content[:1000] if len(content) > 1000 else content
    lines = content.split('\n')
    first_lines = '\n'.join(lines[:10])
    
    print("❌ All parsing strategies failed")
    print(f"Content starts with: {content[:100]}")
    print(f"Content ends with: {content[-100:]}")
    print(f"First 10 lines:\n{first_lines}")
    
    raise ValueError(f"Failed to parse JSON after all strategies. Content length: {len(content)}. Preview: {content_preview}...")


def extract_resume_data_from_response(parsed_response: Dict[str, Any]) -> Dict[str, Any]:
    """Extract resume data from potentially nested API response structure"""
    print("Extracting resume data from response structure...")
    
    # If it's already in the expected format, return as-is
    if 'personalInfo' in parsed_response or 'experience' in parsed_response:
        print("Response is already in resume data format")
        return parsed_response
    
    # Check for nested structures commonly returned by APIs
    possible_paths = [
        ['data', 'content'],      # {data: {content: {resumeData}}}
        ['content'],              # {content: {resumeData}}
        ['data'],                 # {data: {resumeData}}
        ['result'],               # {result: {resumeData}}
        ['response'],             # {response: {resumeData}}
    ]
    
    for path in possible_paths:
        try:
            current = parsed_response
            print(f"Trying path: {' -> '.join(path)}")
            
            for key in path:
                if isinstance(current, dict) and key in current:
                    current = current[key]
                    print(f"Found key '{key}', continuing...")
                else:
                    print(f"Key '{key}' not found, trying next path")
                    break
            else:
                # If we made it through the entire path
                if isinstance(current, dict) and ('personalInfo' in current or 'experience' in current):
                    print(f"Found resume data at path: {' -> '.join(path)}")
                    return current
                else:
                    print("Path complete but doesn't contain resume data")
        except Exception as e:
            print(f"Error following path {path}: {str(e)}")
            continue
    
    # Return original response if we can't extract resume data
    print("Could not extract resume data, returning original response")
    return parsed_response

def call_gemini_api_cached(resume_data: str, model: str = "gemini-2.5-flash-lite") -> Dict[str, Any]:
    """Higher-level function that returns analysis results and debug information"""
    try:
        # Build dynamic prompt
        dynamic_prompt = build_dynamic_prompt(resume_data)
        
        # Count input tokens
        input_tokens = count_tokens(dynamic_prompt, model)
        
        # Call API with caching
        analysis_content = call_gemini_with_cache_and_retry(dynamic_prompt, model, 2)
        
        # Count output tokens
        output_tokens = count_tokens(analysis_content, model)
        
        # Calculate cost
        total_cost = calculate_total_cost(input_tokens, output_tokens)
        
        # Get cache expiry info
        now = time.time() * 1000
        left_time_minutes = 0
        if resume_cache_expiry:
            left_time_minutes = max(round((resume_cache_expiry - now) / 60000), 0)
        
        return {
            "analysis_content": analysis_content,
            "debug_info": {
                "model": model,
                "input_tokens": input_tokens,
                "output_tokens": output_tokens,
                "total_cost": total_cost,
                "prompt_length": len(dynamic_prompt),
                "resume_length": len(resume_data),
                "response_length": len(analysis_content),
                "timestamp": time.strftime("%Y-%m-%dT%H:%M:%S.%fZ"),
                "resume_cache_id": resume_cache_id,
                "left_time_in_expiry_minutes": left_time_minutes,
            }
        }
        
    except Exception as e:
        print(f"Cached analysis failed: {str(e)}")
        raise e

def transform_text_to_resume_data_cached(raw_text: str) -> dict:
    """Transform text to structured data with Gemini API prompt caching"""
    print(f'Processing raw_text of length: {len(raw_text)}')
    
    # Check regular cache first (for text extraction results)
    text_hash = _get_text_hash(raw_text)
    cache_key = f"gemini_transform_cached:{text_hash}"

    cached_result = _get_from_cache(cache_key)
    if cached_result is not None:
        print("Using cached result")
        return cached_result
        
    try:
        # Use cached API call
        print("Calling Gemini API with prompt caching...")
        api_response = call_gemini_api_cached(raw_text)
        
        # Extract the actual content from the response
        content = api_response.get("analysis_content")
        print(f"Received content of length: {content}")
        debug_info = api_response.get("debug_info", {})

        # Log debug information
        print(f"API Stats: {debug_info}")

        if not content:
            error_result = {"error": "Gemini API returned empty content"}
            return error_result

        # Try to parse JSON from the response
        try:
            # First try to extract JSON from code blocks (like your Node.js code)
            import re
            json_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', content)
            if json_match:
                print("Found JSON in code block format")
                result = json.loads(json_match.group(1))
            else:
                print("Attempting to parse entire response as JSON")
                result = json.loads(content)

            # Cache successful result
            _set_cache(cache_key, result)
            return result

        except json.JSONDecodeError as e:
            print(f"JSON Parse Error: {str(e)}")
            print(f"Content that failed to parse: {content[:500]}...")
            error_result = {"error": f"Failed to parse JSON response: {str(e)}"}
            # Don't cache JSON parsing errors
            return error_result

    except Exception as e:
        print(f"Gemini API Exception: {str(e)}")
        return {
            "success": False,
            "error": f"Gemini API error: {str(e)}",
            "error_type": type(e).__name__,
            "debug": {"cache_hit": False}
        }

def parse_resume_cached(filename: str, content: bytes) -> dict:
    """Main parsing function with prompt caching - enhanced version"""
    # Check for complete cached result first
    file_hash = _get_file_hash(content)
    cache_key = f"full_parse_cached:{file_hash}"

    cached_result = _get_from_cache(cache_key)
    if cached_result is not None:
        print("Using cached complete parsing result")
        return cached_result

    # Process text extraction (with its own caching)
    if filename == 'resume.txt':
        raw_text = content.decode('utf-8') if isinstance(content, bytes) else content
    else:
        raw_text = extract_text_from_resume(filename, content)
    
    # Use cached API call for structured data transformation
    structured_data = transform_text_to_resume_data_cached(raw_text)

    # Cache the complete result
    _set_cache(cache_key, structured_data)
    return structured_data

